# processors/exporter.py
import json
import os
from typing import List, Dict, Any
from datetime import datetime

# Проверка доступности библиотек
try:
    import docx
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.units import cm
    PDF_OUTPUT_AVAILABLE = True
except ImportError:
    PDF_OUTPUT_AVAILABLE = False

class ResultExporter:
    """Класс для экспорта результатов в различные форматы"""
    
    def __init__(self):
        self.download_dir = None
        self.font_name = None
        self._register_fonts()
    
    def _register_fonts(self):
        """Регистрация шрифтов для поддержки кириллицы"""
        if not PDF_OUTPUT_AVAILABLE:
            return
            
        try:
            # Попробуем зарегистрировать системные шрифты
            import platform
            system = platform.system()
            
            if system == "Windows":
                # Попробуем найти Arial Unicode или другие шрифты с поддержкой кириллицы
                font_paths = [
                    r"C:\Windows\Fonts\arialuni.ttf",
                    r"C:\Windows\Fonts\arial.ttf", 
                    r"C:\Windows\Fonts\times.ttf",
                    r"C:\Windows\Fonts\calibri.ttf"
                ]
            elif system == "Darwin":  # macOS
                font_paths = [
                    "/System/Library/Fonts/Arial Unicode.ttf",
                    "/System/Library/Fonts/Arial.ttf",
                    "/System/Library/Fonts/Times New Roman.ttf"
                ]
            else:  # Linux
                font_paths = [
                    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
                    "/usr/share/fonts/TTF/arial.ttf"
                ]
            
            # Попробуем зарегистрировать первый доступный шрифт
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('CustomFont', font_path))
                        self.font_name = 'CustomFont'
                        print(f"✅ Шрифт зарегистрирован: {font_path}")
                        return
                    except Exception as e:
                        print(f"❌ Ошибка регистрации шрифта {font_path}: {e}")
                        continue
            
            # Если системные шрифты не найдены, используем встроенный шрифт
            self.font_name = 'Helvetica'
            print("ℹ️ Используется стандартный шрифт Helvetica")
            
        except Exception as e:
            self.font_name = 'Helvetica'
            print(f"ℹ️ Используется стандартный шрифт из-за ошибки: {e}")
    
    def set_download_directory(self, download_dir: str):
        """Установка директории для сохранения файлов"""
        self.download_dir = download_dir
        os.makedirs(download_dir, exist_ok=True)
    
    def export_to_json(self, chunks: List[Dict], filename: str, for_rag: bool = False) -> str:
        """Экспорт в JSON формат"""
        if not self.download_dir:
            raise ValueError("Директория для скачивания не установлена")
        
        # Создаем имя файла
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{filename}_chunks_{timestamp}.json"
        output_path = os.path.join(self.download_dir, output_filename)
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                if for_rag:
                    # Для RAG - чистый JSON
                    json.dump(chunks, f, ensure_ascii=False, indent=2)
                else:
                    # Для чтения - форматированный вывод
                    formatted_data = []
                    for i, chunk in enumerate(chunks):
                        formatted_chunk = {
                            "id": i + 1,
                            "category": chunk.get('category', chunk.get('metadata', {}).get('category', 'N/A')),
                            "content": chunk.get('content', chunk.get('text', '')),
                            "metadata": chunk.get('metadata', {})
                        }
                        formatted_data.append(formatted_chunk)
                    json.dump(formatted_data, f, ensure_ascii=False, indent=2)
            
            return output_path
        except Exception as e:
            raise Exception(f"Ошибка экспорта в JSON: {str(e)}")
    
    def export_to_txt(self, chunks: List[Dict], filename: str, for_rag: bool = False) -> str:
        """Экспорт в TXT формат"""
        if not self.download_dir:
            raise ValueError("Директория для скачивания не установлена")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{filename}_chunks_{timestamp}.txt"
        output_path = os.path.join(self.download_dir, output_filename)
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                if for_rag:
                    # Для RAG - JSON в одну строку
                    json.dump(chunks, f, ensure_ascii=False)
                else:
                    # Для чтения - форматированный текст
                    for i, chunk in enumerate(chunks):
                        category = chunk.get('category', chunk.get('metadata', {}).get('category', 'N/A'))
                        content = chunk.get('content', chunk.get('text', ''))
                        f.write(f"### ЧАНК #{i+1} | {category}\n{content}\n\n")
            
            return output_path
        except Exception as e:
            raise Exception(f"Ошибка экспорта в TXT: {str(e)}")
    
    def export_to_docx(self, chunks: List[Dict], filename: str, for_rag: bool = False) -> str:
        """Экспорт в DOCX формат"""
        if not DOCX_AVAILABLE:
            raise Exception("Библиотека python-docx не установлена")
        
        if not self.download_dir:
            raise ValueError("Директория для скачивания не установлена")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{filename}_chunks_{timestamp}.docx"
        output_path = os.path.join(self.download_dir, output_filename)
        
        try:
            doc = docx.Document()
            
            if for_rag:
                # Для RAG - чистый JSON
                p = doc.add_paragraph(json.dumps(chunks, ensure_ascii=False, indent=2))
                p.style.font.name = 'Consolas'
                if hasattr(p.style.font, 'size'):
                    p.style.font.size = Pt(8)
            else:
                # Для чтения - форматированный документ
                doc.add_heading('Отчет Чанкер РА', 0)
                doc.add_paragraph(f'Всего чанков: {len(chunks)}')
                doc.add_paragraph(f'Дата обработки: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
                doc.add_paragraph('_' * 50)
                
                for i, chunk in enumerate(chunks):
                    if i % 100 == 0 and i > 0:
                        # Добавляем разрыв страницы каждые 100 чанков
                        doc.add_page_break()
                    
                    category = chunk.get('category', chunk.get('metadata', {}).get('category', 'N/A'))
                    content = chunk.get('content', chunk.get('text', ''))
                    
                    # Добавляем заголовок чанка
                    heading = doc.add_heading(f'Чанк #{i+1}: {category}', level=1)
                    if hasattr(heading.style.font, 'color'):
                        from docx.shared import RGBColor
                        heading.style.font.color.rgb = RGBColor(33, 150, 243)
                    
                    # Добавляем содержимое
                    doc.add_paragraph(content)
                    
                    if i < len(chunks) - 1:
                        doc.add_paragraph('_' * 30)
            
            doc.save(output_path)
            return output_path
        except Exception as e:
            raise Exception(f"Ошибка экспорта в DOCX: {str(e)}")
    
    def export_to_pdf(self, chunks: List[Dict], filename: str, for_rag: bool = False) -> str:
        """Экспорт в PDF формат с поддержкой кириллицы"""
        if not PDF_OUTPUT_AVAILABLE:
            raise Exception("Библиотека reportlab не установлена")
        
        if not self.download_dir:
            raise ValueError("Директория для скачивания не установлена")
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{filename}_chunks_{timestamp}.pdf"
        output_path = os.path.join(self.download_dir, output_filename)
        
        try:
            doc = SimpleDocTemplate(output_path, pagesize=A4, 
                                  rightMargin=1*cm, leftMargin=1*cm, 
                                  topMargin=1*cm, bottomMargin=1*cm)
            story = []
            styles = getSampleStyleSheet()
            
            # Создаем стили с поддержкой кириллицы
            if self.font_name:
                title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], 
                                           fontName=self.font_name, fontSize=14, 
                                           spaceAfter=10, textColor=colors.darkblue)
                normal_style = ParagraphStyle('NormalStyle', parent=styles['Normal'], 
                                            fontName=self.font_name, fontSize=8, leading=10)
                code_style = ParagraphStyle('CodeStyle', parent=styles['Normal'], 
                                          fontName=self.font_name, fontSize=7, leading=9, 
                                          textColor=colors.black)
            else:
                title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], 
                                           fontSize=14, spaceAfter=10, textColor=colors.darkblue)
                normal_style = ParagraphStyle('NormalStyle', parent=styles['Normal'], 
                                            fontSize=8, leading=10)
                code_style = ParagraphStyle('CodeStyle', parent=styles['Normal'], 
                                          fontSize=7, leading=9, textColor=colors.black)
            
            if for_rag:
                # Для RAG - чистый JSON
                json_str = json.dumps(chunks, ensure_ascii=False, indent=2)
                lines = json_str.split('\n')
                
                for i, line in enumerate(lines):
                    # Экранируем специальные символы HTML
                    safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(safe_line, code_style))
            else:
                # Для чтения - форматированный PDF
                story.append(Paragraph('Отчет Чанкер РА', title_style))
                story.append(Paragraph(f'Чанков: {len(chunks)}', normal_style))
                story.append(Paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}', normal_style))
                story.append(Spacer(1, 10))
                
                data = [["ID", "Категория", "Текст"]]
                
                for i, chunk in enumerate(chunks[:1000]):  # Ограничиваем 1000 чанков для производительности
                    content = chunk.get('content', chunk.get('text', ''))
                    # Экранируем специальные символы и заменяем переносы строк
                    content_safe = content.replace('\n', '<br/>').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    p_content = Paragraph(content_safe, normal_style)
                    cat_val = chunk.get('category', chunk.get('metadata', {}).get('category', 'N/A'))
                    data.append([str(i+1), cat_val, p_content])
                
                table = Table(data, colWidths=[1.5*cm, 4*cm, 12.5*cm])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), self.font_name if self.font_name else 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, 1), (-1, -1), self.font_name if self.font_name else 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 7),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(table)
            
            doc.build(story)
            return output_path
        except Exception as e:
            raise Exception(f"Ошибка экспорта в PDF: {str(e)}")

# Глобальный экземпляр экспортера
exporter = ResultExporter()
